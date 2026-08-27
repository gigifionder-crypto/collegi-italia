# -*- coding: utf-8 -*-
import re, importlib.util
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load(mod, path):
    spec = importlib.util.spec_from_file_location(mod, path)
    m = importlib.util.module_from_spec(spec); spec.loader.exec_module(m)
    return m.PARA

PARA = load("c01","/home/claude/c01.py") + load("c02","/home/claude/c02.py") + load("c03","/home/claude/c03.py")

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Twips(1440)

st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def force(run, bold=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(12); run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rf.set(qn(a), "Times New Roman")

def add(kind, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.15
    pf.space_before = Pt(0); pf.space_after = Pt(6)
    if kind == "T":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); force(r, bold=True)
    elif kind == "C":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); force(r, bold=False)
    elif kind == "H":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(12)
        r = p.add_run(text); force(r, bold=True)
    elif kind == "SP":
        r = p.add_run(""); force(r)
    elif kind == "PB":
        r = p.add_run(); force(r); r.add_break(WD_BREAK.PAGE)
    else:  # P
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.75)
        for i, seg in enumerate(re.split(r"\*\*", text)):
            if seg == "": continue
            r = p.add_run(seg); force(r, bold=(i % 2 == 1))

for kind, text in PARA:
    add(kind, text)

# Zoom al cento per cento
zoom = OxmlElement("w:zoom"); zoom.set(qn("w:percent"), "100")
doc.settings.element.insert(0, zoom)

out = "/mnt/user-data/outputs/ITALIA_NERA_PROMPT_EVOLUTO_V65_METODO_DE_MICHELE_E_CAMERA_ISTRUTTORIA_D.docx"
doc.save(out)
print("SALVATO:", out, "| blocchi:", len(PARA))
